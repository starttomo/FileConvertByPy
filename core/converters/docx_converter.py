# core/converters/docx_converter.py
from docx import Document
from pathlib import Path
from core.converters.base_converter import BaseConverter
from core.exceptions import FileReadError, FileWriteError, UnsupportedFormatError, FileConversionError
from core.registry import register_converter
import pypandoc
class DocxConverter(BaseConverter):
    """处理DOCX文件的转换"""

    @classmethod
    def supported_formats(cls) -> dict:
        return {
            'docx': ['txt', 'html', 'pdf'],  # 添加 'pdf' 支持
            # 可以添加更多支持的输出格式
        }

    @classmethod
    def convert(cls, input_path: str, output_path: str, input_ext: str, output_ext: str):
        try:
            cls._ensure_output_dir_exists(output_path)

            if output_ext == 'txt':
                cls._convert_to_txt(input_path, output_path)
            elif output_ext == 'html':
                cls._convert_to_html(input_path, output_path)
            elif output_ext == 'pdf':  # 添加对 pdf 转换的处理
                cls._convert_to_pdf(input_path, output_path)
            else:
                raise UnsupportedFormatError(f"不支持将 docx 转换为 {output_ext}")
        except Exception as e:
            raise FileConversionError(f"DOCX转换失败: {str(e)}")

    @classmethod
    def _convert_to_txt(cls, input_path: str, output_path: str):
        try:
            doc = Document(input_path)
            with open(output_path, 'w', encoding='utf-8') as f:
                for para in doc.paragraphs:
                    f.write(para.text + '\n')
        except Exception as e:
            raise FileReadError(f"读取DOCX文件失败: {str(e)}")

    @classmethod
    def _convert_to_html(cls, input_path: str, output_path: str):
        try:
            doc = Document(input_path)
            html_content = ['<!DOCTYPE html><html><head><meta charset="UTF-8"></head><body>']

            for para in doc.paragraphs:
                html_content.append(f'<p>{para.text}</p>')

            html_content.append('</body></html>')

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(html_content))
        except Exception as e:
            raise FileWriteError(f"写入HTML文件失败: {str(e)}")

    @classmethod
    def _convert_to_pdf(cls, input_path: str, output_path: str):
        try:
            # 使用xelatex引擎并添加中文支持
            pypandoc.convert_file(
                input_path,
                'pdf',
                outputfile=output_path,
                format='docx',
                extra_args=[
                    '--pdf-engine=xelatex',  # 使用xelatex引擎支持中文
                    '-V', 'mainfont=SimHei',  # 指定中文字体
                    '-V', 'geometry:a4paper',  # 设置纸张大小
                ]
            )
        except ImportError:
            raise FileConversionError("PDF转换需要安装pypandoc和系统上的LaTeX发行版")
        except Exception as e:
            raise FileWriteError(f"生成PDF失败: {str(e)}")

# 注册转换器
for input_ext, output_exts in DocxConverter.supported_formats().items():
    for output_ext in output_exts:
        register_converter(input_ext, output_ext, DocxConverter)